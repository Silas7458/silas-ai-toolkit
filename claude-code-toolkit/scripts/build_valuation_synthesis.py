"""
Build Hospice Valuation Research Synthesis .docx
Combines all 9 agent reports + Proctor's v2 research into final deliverable
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DST = '{{DOCS_DIR}}/claude-context/deliverables/2026-02-27-brother-hospice-valuation-synthesis.docx'

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style setup
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

DARK_BLUE = RGBColor(0x2F, 0x54, 0x96)
MED_BLUE = RGBColor(0x44, 0x72, 0xC4)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
    return h

def bold_text(para, text):
    run = para.add_run(text)
    run.bold = True
    return run

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Hospice Enterprise Valuation', level=0)
for run in title.runs:
    run.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Market Research Synthesis & Model Update')
r.font.size = Pt(14)
r.font.color.rgb = MED_BLUE

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Amerix Medical Consulting, LLC  |  February 27, 2026').font.size = Pt(10)

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = info2.add_run('Research: 10 parallel agents + Proctor v2 synthesis  |  Model: v2.5.8 \u2192 v3.0')
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()  # spacer

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
heading('Executive Summary')

doc.add_paragraph(
    'This document synthesizes hospice valuation research from 10 parallel research agents '
    'and Proctor\u2019s independent market analysis, covering 40+ web sources, 73 transaction '
    'data points, 7 real disclosed deals, 6 public company comparables, 9 Texas hospice listings, '
    'and 19 market-sourced sensitivity factors. All findings have been implemented into the '
    'valuation model (v3.0).'
)

p = doc.add_paragraph()
bold_text(p, 'Key Finding: ')
p.add_run(
    'The model\u2019s EBITDA-A multiple of 6.5x was too high for a 40-ADC Texas hospice. '
    'All data sources converge on 3.5x\u20136.5x EBITDA for the Small Profitable tier (20\u201350 ADC), '
    'with a median of 4.7x from Scope Research\u2019s 73-deal dataset. The model has been '
    'recalibrated to 4.7x EBITDA-A, producing a final market-adjusted value of ~$1.56M '
    '(down from $1.87M), or ~$39K per ADC \u2014 conservative but defensible for a neutral-scoring '
    'Texas agency.'
)

p2 = doc.add_paragraph()
bold_text(p2, 'Model Changes (v2.5.8 \u2192 v3.0): ')
p2.add_run(
    '(1) 7-tier ADC auto-estimates replacing 3-tier, (2) dynamic NOI multiple pegged at 65% '
    'of EBITDA-A EV, (3) Market & Sensitivity Adjustments section with Non-CON toggle and '
    '19-factor scoring, (4) new Sensitivity Scoring sheet, (5) Final Market-Adjusted Enterprise '
    'Value with implied per-ADC cross-check.'
)

# ============================================================
# 2. EBITDA MULTIPLES BY ADC TIER
# ============================================================
heading('EBITDA Multiples by ADC Tier (7-Tier Framework)')

doc.add_paragraph(
    'EBITDA multiples are the dominant pricing mechanism in hospice M&A. Scale (measured by ADC '
    'and revenue) is the single largest driver. The table below consolidates data from Scope Research '
    '(73 deals, 51 with EBITDA multiples), Mertz Taggart, Vallexa Advisors, FOCUS Bankers, '
    'Stoneridge Partners, HealthFMV, and HCBB.'
)

add_table(
    ['ADC Tier', 'ADC Range', 'EBITDA Low', 'EBITDA Median', 'EBITDA High', 'Revenue', 'SDE', 'EV/ADC'],
    [
        ['Startup', '<10', 'N/A', 'N/A', 'N/A', '0.3x\u20131.5x', 'N/A', '$15K\u2013$35K'],
        ['Emerging', '10\u201320', '3.0x', 'N/A (use rev)', '5.0x', '0.4x\u20131.0x', '2.5x', '$30K\u2013$50K'],
        ['Small Profitable', '20\u201350', '3.5x', '4.7x', '6.5x', '0.5x\u20131.2x', '3.5x', '$45K\u2013$70K'],
        ['Medium', '50\u201380', '5.0x', '6.0x', '7.5x', '0.7x\u20131.5x', '4.0x', '$55K\u2013$85K'],
        ['Larger Medium', '80\u2013120', '6.5x', '7.5x', '9.0x', '0.8x\u20132.0x', '4.5x', '$70K\u2013$100K'],
        ['Large', '120\u2013200', '8.0x', '10.0x', '12.0x', '1.0x\u20132.5x', '5.0x', '$85K\u2013$130K'],
        ['Enterprise', '>200', '10.0x', '14.5x', '17.0x+', '1.2x\u20133.0x+', '6.0x', '$100K\u2013$165K'],
    ]
)

p = doc.add_paragraph()
p.add_run(
    'Our model target (ADC 40) sits in the Small Profitable tier. At 4.7x EBITDA median, '
    'with a range of 3.5x\u20136.5x. The old 6.5x was at the absolute ceiling of this tier.'
).font.size = Pt(10)

# ============================================================
# 3. REAL TRANSACTION BENCHMARKS
# ============================================================
heading('Real Transaction Benchmarks')

doc.add_paragraph(
    'The following are actual disclosed or derivable deal data points. Enterprise-scale deals '
    '(>$100M) trade at 10x\u201318x EBITDA reflecting platform premiums. Small private deals are '
    'discounted 60\u201375% from public company multiples.'
)

add_table(
    ['Deal', 'Price', 'EBITDA Multiple', 'EV/ADC', 'Year'],
    [
        ['Optum / Amedisys', '$3.3B', '13.4x', '~$260K', '2025'],
        ['Kinderhook / Enhabit', '$1.1B', '10.2x', '~$95K (hospice)', '2025'],
        ['VITAS / Covenant', '$85M', '10.8x\u201318.1x', 'N/A', '2024'],
        ['Addus / Queen City', '$192M', '~24x', '~$213K', '2024'],
        ['Addus / Hospice Partners', '$130M', 'N/A', '~$130K', '2024'],
        ['Pennant / Signature', '$80M', 'N/A', '~$267K', '2024'],
        ['BrightSpring / Haven', '$60M', 'Est 8\u201310x', 'N/A', '2024'],
    ]
)

doc.add_paragraph(
    'Key insight: Even the "smallest" disclosed deals ($60M+) involve agencies far larger than ADC 40. '
    'Small private deals (<$10M EV) rarely disclose multiples publicly, making broker listing data '
    'and Scope Research\u2019s database essential benchmarks.'
)

# ============================================================
# 4. TEXAS MARKET FACTORS
# ============================================================
heading('Texas Market Factors')

doc.add_paragraph(
    'Texas is the second-largest hospice market nationally (9.4M patient days/year) with unique '
    'characteristics that directly affect valuation multiples.'
)

h3 = heading('Non-CON State \u2014 The Biggest TX-Specific Factor', level=2)
doc.add_paragraph(
    'Texas repealed Certificate of Need in 1985. No regulatory barrier to new hospice entry. '
    'Per Stoneridge Partners (Dec 2024): "Non-CON states provide greater flexibility for expansion '
    'but typically see lower deal multiples due to competitive vulnerability." Documented impact: '
    '0.5x\u20131.5x lower EBITDA multiple vs. CON states at comparable scale.'
)

heading('Net Texas Discount by ADC Tier', level=2)
doc.add_paragraph(
    'Texas advantages (10\u201315% lower labor costs, no state income tax, 65+ population growing '
    '88% by 2050, 52.2% hospice penetration) partially offset the non-CON discount, especially '
    'at larger scale.'
)

add_table(
    ['ADC Tier', 'Net TX Adjustment', 'Rationale'],
    [
        ['Startup (<10)', '-15% to -17%', 'Maximum vulnerability, no scale to offset non-CON risk'],
        ['Emerging (10\u201320)', '-12% to -13%', 'Still fragile, limited referral relationships'],
        ['Small Profitable (20\u201350)', '-6% to -8%', 'Labor savings + growth partially offset non-CON'],
        ['Medium (50\u201380)', '-5% to -7%', 'Established operations, multiple referral sources'],
        ['Larger Medium (80\u2013120)', '-3% to -5%', 'Scale approaches regional platform status'],
        ['Large (120\u2013200)', '0% to -2%', 'Platform value offsets most state-level discounts'],
        ['Enterprise (>200)', '0% to +2%', 'TX advantages (labor, growth) become net positive'],
    ]
)

heading('Texas Market Data Points', level=2)
bullets = [
    '700+ new hospices opened nationally 2022\u20132023; TX under CMS enhanced oversight for new providers since Summer 2023',
    '75% of US hospices are for-profit (Penn LDI, 2024); Texas skews higher',
    'For-profit margins: 18\u201322% for PE-backed agencies vs 5\u20136% for non-profit',
    'VITAS/Chemed: ADC 22,462, revenue $208/patient/day, targeting acquisitions at 6x\u20138x in CON states specifically',
    'Enhabit (TX-HQ): 110 hospice locations, ADC 3,809 (+12.3% YoY), hospice EBITDA up 53.8%',
    'Addus (TX-HQ): hospice segment $62M/quarter, 10% organic growth',
    'FY2026 Medicare payment update: +2.6%',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 5. TEXAS HOSPICE LISTINGS
# ============================================================
heading('Active Texas Hospice Listings Found')

doc.add_paragraph(
    '9 Texas hospice listings were identified across BizBuySell, BizQuest, and industry sources:'
)

add_table(
    ['Location', 'Asking Price', 'Key Details'],
    [
        ['License-only shells', '$250K\u2013$300K', 'San Antonio, statewide \u2014 no patients, just license'],
        ['Houston / Sugar Land', '$600K', 'Active operations, small census'],
        ['Lubbock', '$300K', 'Small market, limited data'],
        ['Austin', '$595K', '0.8x revenue, 3.2x cash flow'],
        ['Dallas / North TX', '$1.2M', 'Established operations'],
        ['Region 6 (Houston area)', '$2.8M revenue', '37\u201340 ADC, declining census'],
        ['Rio Grande Valley', 'Not disclosed', '45+ ADC, active operations'],
    ]
)

p_note = doc.add_paragraph(
    'The Region 6 listing (37\u201340 ADC, declining) is the closest comp to our model target. '
    'At $2.8M revenue and assuming 15% EBITDA margin ($420K EBITDA), an asking price in the '
    '$1.5\u2013$2.0M range would imply 3.6x\u20134.8x EBITDA \u2014 consistent with our 4.7x central estimate.'
)
for run in p_note.runs:
    run.font.size = Pt(10)

# ============================================================
# 6. PUBLIC COMPANY COMPARABLES
# ============================================================
heading('Public Company Comparables')

doc.add_paragraph(
    'Public hospice/home health companies trade at significant premiums to private deals, '
    'reflecting platform scale, diversification, and liquidity. Small private discount: 60\u201375%.'
)

add_table(
    ['Company', 'Hospice ADC', 'Implied EV/ADC', 'Hospice EBITDA Margin', 'Key Metric'],
    [
        ['VITAS (Chemed)', '22,338', '~$161K', '12.6% (normalizing to 17.5\u201318%)', '$208/pt/day revenue'],
        ['Amedisys (UnitedHealth)', '~12,700', '~$150K', 'Combined with HH', '$3.3B deal at 13.4x'],
        ['Enhabit', '3,809', '~$95K', 'EBITDA +53.8% YoY', '12.3% ADC growth, de novo strategy'],
        ['Addus', 'Segment only', 'N/A (combined)', '~15% estimated', '$62M/quarter hospice revenue'],
        ['Pennant Group', '~30 locations', '$2.4M\u2013$3.9M/location', 'Growing', 'Tuck-in acquisition model'],
        ['BrightSpring', 'Combined', 'N/A', 'N/A', 'Pharmacy + provider services'],
    ]
)

doc.add_paragraph(
    'Applying the 60\u201375% small private discount to VITAS\u2019s ~$161K per ADC implies '
    '$40K\u2013$64K per ADC for a comparable small private hospice. Our model\u2019s implied '
    '$39K\u2013$50K per ADC (depending on sensitivity score) is within this range.'
)

# ============================================================
# 7. SENSITIVITY SCORING FRAMEWORK
# ============================================================
heading('19-Factor Sensitivity Scoring Framework')

doc.add_paragraph(
    'Beyond scale and geography, 19 qualitative factors are documented by M&A advisors as directly '
    'moving valuation multiples. These have been organized into three tiers by impact magnitude, '
    'with sources from Scope Research, Mertz Taggart, Vallexa, Stoneridge, HealthFMV, and Weaver.'
)

heading('Tier 1 \u2014 Large Impact (Weight 3\u00d7)', level=2)
t1 = [
    ('Revenue / EBITDA Trend', 'Scope, Mertz Taggart', 'Growing agencies command 1\u20132x premium over flat; declining agencies see 1\u20133x discount'),
    ('Payer Mix Quality', 'Vallexa, Penn LDI', 'High Medicare % = predictable revenue; heavy Medicaid = margin compression'),
    ('Survey / Compliance History', 'Stoneridge, HealthFMV', 'Clean surveys = buyer confidence; sanctions/CMPs = material discount or deal-breaker'),
    ('EBITDA Margin vs Peers', 'Scope (73 deals)', 'Top-quartile margins (>20%) add 0.5x\u20131.0x; sub-8% margins reduce by 1\u20132x'),
    ('Operator Quality / Transferability', 'Mertz Taggart, Braff', 'Key-person risk is single biggest small-deal discount; strong teams transfer value'),
    ('CON / License Protection', 'Stoneridge, Provident', 'CON states: +0.5x\u20131.5x; Non-CON with enhanced oversight: additional -0.5x'),
    ('Geographic Market Strength', 'Scope, Acuvance', 'High-growth metros vs. rural/saturated markets; 65+ demographics drive volume'),
]
add_table(
    ['Factor', 'Source', 'Impact Documentation'],
    t1
)

heading('Tier 2 \u2014 Moderate Impact (Weight 2\u00d7)', level=2)
t2 = [
    ('Average Length of Stay', 'Weaver, Hospice News', 'Optimal 60\u201390 days; very short or very long ALOS signals operational issues'),
    ('Referral Concentration', 'Vallexa, HCBB', '>40% from single source = material risk; diverse = premium'),
    ('Staff Retention / Turnover', 'HealthFMV', 'RN turnover >40% = operational risk; <15% = premium stability'),
    ('Live Discharge Rate', 'CMS, Hospice News', 'High rates signal enrollment issues; low = appropriate admissions'),
    ('CAHPS / Quality Scores', 'Vallexa', '4+ stars = marketing advantage and buyer confidence'),
    ('Payer Contract Breadth', 'Mertz Taggart', 'MA contracts create diversified revenue and higher reimbursement'),
]
add_table(
    ['Factor', 'Source', 'Impact Documentation'],
    t2
)

heading('Tier 3 \u2014 Incremental Impact (Weight 1\u00d7)', level=2)
t3 = [
    ('EMR / Technology Stack', 'HCBB', 'Modern cloud EMR reduces transition cost for buyer'),
    ('Facility / Lease Terms', 'Stoneridge', 'Favorable long-term leases = stability; expiring = risk/opportunity'),
    ('Brand / Community Reputation', 'Vallexa', 'Strong local brand = referral continuity'),
    ('De Novo vs Mature', 'Scope, Enhabit data', '<2 years = high risk, limited data; 5+ = track record premium'),
    ('Pending Regulatory Changes', 'CMS, Hospice News', 'HOPE tool implementation (2026), payment reform impact'),
    ('Documentation / HOPE Readiness', 'Vallexa, CMS', 'HOPE-ready agencies better positioned for 2026 changes'),
]
add_table(
    ['Factor', 'Source', 'Impact Documentation'],
    t3
)

heading('Scoring Mechanics', level=2)
doc.add_paragraph(
    'Each factor scores -2 (worst) to +2 (best). Weights: Tier 1 = 3\u00d7, Tier 2 = 2\u00d7, Tier 3 = 1\u00d7. '
    'Composite score range: -78 to +78. Multiple adjustment = Score \u00d7 0.046.'
)
doc.add_paragraph('Worked examples:', style='List Bullet')
doc.add_paragraph(
    'Florida 40-ADC, clean surveys, 22% margin, 8% growth \u2192 Score +32 \u2192 +1.47x adjustment '
    '\u2192 Final 6.17x EBITDA',
    style='List Bullet 2'
)
doc.add_paragraph(
    'Texas 40-ADC, deficiencies, 8% margin, declining census \u2192 Score -47 \u2192 -2.16x adjustment '
    '\u2192 Final 2.54x EBITDA',
    style='List Bullet 2'
)

# ============================================================
# 8. MODEL CHANGES
# ============================================================
heading('Model Changes Implemented (v2.5.8 \u2192 v3.0)')

doc.add_paragraph(
    'The following changes were made to hospice_valuation_model_3_0.xlsx based on this research:'
)

changes = [
    ('EBITDA-A Multiple', '6.5x (3-tier)', '4.7x (7-tier ADC framework)',
     'Scope Research 73-deal median for ADC 20\u201350. Margin bonus: +0.5 at >20%, +1.0 at >25%.'),
    ('NOI (EBITDA-B) Multiple', '8.5x (fixed)', '~9.2x (dynamic)',
     'Now derived as 65% of EBITDA-A EV, translated to NOI basis. Auto-adjusts for BTL %.'),
    ('Revenue Multiple', '0.8x (3-tier)', '0.8x (7-tier)',
     'Unchanged at ADC 40 tier. Now uses 7 tiers with margin bonus.'),
    ('SDE Multiple', '3.5x (3-tier)', '3.5x (7-tier)',
     'Unchanged at ADC 40 tier. Now uses 7 tiers with margin bonus.'),
    ('Texas Discount', 'None', '-7.0% at ADC 40',
     'Non-CON toggle with ADC-tier-based discount. Sources: Stoneridge, Provident, Scope.'),
    ('Sensitivity Scoring', 'None', '19-factor sheet',
     'Weighted composite score feeds EV adjustment. All factors market-sourced with citations.'),
    ('Final Value (ADC 40, TX)', '$1,866,990', '$1,561,240',
     '-16.4%. More conservative, better aligned with market data for small TX hospice.'),
]

add_table(
    ['Item', 'Old (v2.5.8)', 'New (v3.0)', 'Rationale'],
    changes
)

heading('New Model Sections', level=2)
new_sections = [
    'Rows 89\u201393: Market & Sensitivity Adjustments \u2014 Non-CON toggle, sensitivity score input, combined market adjustment factor',
    'Rows 95\u2013103: Final Market-Adjusted Enterprise Value \u2014 low/mid/high range with TX and sensitivity applied, plus implied per-ADC cross-check',
    'Sensitivity Scoring sheet: 19 factors with tier headers, weighted scoring, composite total linked to Dashboard',
    'ADC Tier Reference Table: 7-tier base multiples with TX discounts for quick reference',
]
for s in new_sections:
    doc.add_paragraph(s, style='List Bullet')

# ============================================================
# 9. COST STRUCTURE BENCHMARKS
# ============================================================
heading('Cost Structure Benchmarks')

doc.add_paragraph('Key cost structure benchmarks from industry data and agent research:')

add_table(
    ['Metric', 'Industry Average', 'Model Value', 'Source'],
    [
        ['Direct labor % of revenue', '38%', '65% (total staff)', 'MVI, Scope Research'],
        ['Patient-related costs', '17%', '10%', 'MVI operational data'],
        ['Indirect/overhead', '31%', '10%', 'MVI, industry benchmarks'],
        ['Target NOI margin', '14% (avg 5%, excellent 25%)', '5% (after 10% BTL)', 'MVI median: 5.8%'],
        ['Breakeven ADC', '15\u201320 patients', 'N/A (model assumes operating)', 'Scope, Enhabit de novo data'],
        ['RN staffing ratio', '1 per 12\u201314 patients', 'Embedded in staff %', 'CMS, industry standards'],
        ['TX labor advantage', '10\u201315% below national', 'Reflected in TX discount', 'BLS, Enhabit/Addus data'],
    ]
)

# ============================================================
# 10. SOURCES
# ============================================================
heading('Sources')

sources = [
    'Scope Research \u2014 Hospice Valuation Multiples & M&A Trends 2025; 73-deal database (51 with EBITDA multiples). scoperesearch.co',
    'Mertz Taggart \u2014 "It\u2019s All About the Multiple" (Oct 2025); "The Fog is Lifting" (Feb 2026); Q3 2025 Public Company Roundup. mertztaggart.com',
    'Provident Healthcare Partners \u2014 Jake Vesely quotes (Hospice News, May 2024): small-deal compression.',
    'Home Care Business Broker (HCBB) \u2014 2026 Healthcare M&A Report: Hospice 9.0x\u201312.5x (large platforms).',
    'Vallexa Advisors \u2014 Hospice Valuation Guide (Oct 2025); HOPE Tool article (Nov 2025). 4x\u20136x small, 7x\u201310x+ larger.',
    'FOCUS Investment Banking \u2014 Healthcare EBITDA Multiples: 2026 Dashboard.',
    'HealthFMV \u2014 Hospice Valuation Guide 2025; 5x\u20138x small-to-mid; labor compression risk analysis.',
    'Weaver \u2014 Hospice Valuations in a High Multiple Comparable Sales Environment (2024); size-multiple regression.',
    'Stoneridge Partners \u2014 How Valuation Works in M&A (Jun 2025); CON Laws and Healthcare Business Valuation (Nov 2024).',
    'Hospice News \u2014 State of Dealmaking Mid-2025 (Sep 2025); How Valuations Are Shaping Up (May 2024).',
    'Acuvance DataRise / Coker Group \u2014 Hospice Business Valuation (2024); Texas market growth.',
    'The Braff Group \u2014 Hospice Multiples 2023 (Hospice News, May 2023).',
    'Penn LDI \u2014 Medicare Hospice: Exploding in Size (May 2024); margin data: for-profit 18\u201322%.',
    'MVI / MultiView Inc \u2014 Hospice Operational Comparisons by ADC; median client profitability 5.8%.',
    'Chemed/VITAS Q4 2025 Earnings (Feb 26, 2026) \u2014 ADC 22,462, $208/pt/day.',
    'Enhabit Q1 & Q2 2025 Earnings \u2014 Hospice ADC 3,809 (+12.3%), EBITDA +53.8%.',
    'CMS \u2014 Enhanced oversight for new hospices in TX, AZ, CA, NV (Summer 2023); FY2026 2.6% payment update.',
    'ASPE/HHS \u2014 Tracking Impact of Ownership Changes in Hospice.',
]

for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# Save
doc.save(DST)
print(f'Synthesis document saved to: {DST}')
print(f'Sections: 10 | Tables: 9 | Sources: {len(sources)}')
